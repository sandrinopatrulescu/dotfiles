#!/usr/bin/env python3
import csv
import os
import re
import subprocess
import sys
from datetime import datetime
from decimal import Decimal, ROUND_HALF_DOWN
from enum import Enum
from typing import Dict, List, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm, Inches
from docx.table import _Row, _Cell, _Column, Table

# region defaults


PRICE_PER_STUNDEN = 25.0
PRICE_PER_STUNDEN_PL = 3.0
VAT_RATE = 19

DECIMAL_SEPARATOR = ","

STUNDEN_CSV_COLUMN_DATE = 0
STUNDEN_CSV_COLUMN_KN_NR = 1
STUNDEN_CSV_COLUMN_BAU = 2
STUNDEN_CSV_COLUMN_STUNDEN = 3
STUNDEN_CSV_COLUMN_STUNDEN_PL = 4

BLACK = '000000'


# endregion


def get_input_args():
    if not 1 <= len(sys.argv) - 1 <= 3:
        message = f"Usage: python rechnung_data_from_csv.py <csv_file_path> [<first_rechnung_nr>=1] [<price_per_stunden>={PRICE_PER_STUNDEN}]\n"
        sys.stderr.write(message)
        sys.exit(1)

    def get_argument_or_default(index: int, default: any):
        return sys.argv[index] if len(sys.argv) > index else default

    csv_file_path = sys.argv[1]
    first_rechnung_nr = int(get_argument_or_default(2, 1))
    price_per_stunden = float(get_argument_or_default(3, PRICE_PER_STUNDEN))
    return csv_file_path, first_rechnung_nr, price_per_stunden


def date_str_to_date(date_str: str):
    return datetime.strptime(date_str, "%d.%m.%Y")


def is_valid_date(date_str: str):
    try:
        if not bool(re.match(r"^\d{2}\.\d{2}\.\d{4}$", date_str)):
            return False

        date_str_to_date(date_str)
        return True
    except ValueError:
        return False


def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


WorkSessionInfo = Tuple[str, str, float, float]
RechnungInfo = List[WorkSessionInfo]
DateToRechnung = Dict[str, RechnungInfo]


def parse_csv(csv_file_path: str):
    csv_file = open(csv_file_path, mode='r')
    csv_reader = csv.reader(csv_file, delimiter=',')

    date_to_tuple_list: DateToRechnung = {}

    for i, line in enumerate(csv_reader):
        if line[0] == 'datum':
            continue  # skip header
        if not is_valid_date(line[STUNDEN_CSV_COLUMN_DATE]):
            raise ValueError(f"Invalid date {line[STUNDEN_CSV_COLUMN_KN_NR]} at line {i + 1}")
        if not all(map(is_number, [line[STUNDEN_CSV_COLUMN_STUNDEN], line[STUNDEN_CSV_COLUMN_STUNDEN_PL]])):
            raise ValueError(f"Invalid numbers at line {i + 1}")

        date = line[STUNDEN_CSV_COLUMN_DATE]
        kn_nr = line[STUNDEN_CSV_COLUMN_KN_NR]  # ab, auf, um
        bau = line[STUNDEN_CSV_COLUMN_BAU].strip()
        stunden = float(line[STUNDEN_CSV_COLUMN_STUNDEN])
        stunden_pl = float(line[STUNDEN_CSV_COLUMN_STUNDEN_PL])
        tuple_list = date_to_tuple_list.setdefault(date, [])
        tuple_list.append((kn_nr, bau, stunden, stunden_pl))

    return sorted(date_to_tuple_list.items(), key=lambda item: date_str_to_date(item[0]).strftime('%Y.%m.%d'))


def get_issue_date_string(issue_date: datetime):
    return f"München den {issue_date.strftime('%d.%m.%Y')}"


class DocGenerator:
    class PrDataKeys(str, Enum):
        BANK_ACCOUNT = "bank-account"
        HEADER_TEXT_LEFT = "header-text-left"
        HEADER_TEXT_RIGHT = "header-text-right"

    PRICE_TABLE_INFO_COLUMN = 0
    PRICE_TABLE_COMPUTATION_COLUMN = 1
    PRICE_TABLE_RESULT_VALUE_COLUMN = 2
    PRICE_TABLE_RESULT_DETAILS_COLUMN = 3

    def __init__(self):
        self.__pr_data = {}

        pr_dir = os.path.expandvars(os.getenv('OTHERS_M_PR'))
        for pr_data_key in DocGenerator.PrDataKeys:
            key = pr_data_key.value
            self.__pr_data[key] = open(os.path.join(pr_dir, f"{key}.txt"), 'r').read()

    @staticmethod
    def set_table_cell_border(table_cell: _Cell, direction: str, size: int, color=BLACK):
        table_cell_properties = table_cell._tc.get_or_add_tcPr()

        table_cell_borders = table_cell_properties.find(qn('w:tcBorders'))
        if table_cell_borders is None:
            table_cell_borders = OxmlElement('w:tcBorders')
            table_cell_properties.append(table_cell_borders)

        side_border = table_cell_borders.find(qn(f'w:{direction}'))
        if side_border is None:
            side_border = OxmlElement(f'w:{direction}')
            table_cell_borders.append(side_border)

        side_border.set(qn('w:val'), 'single')
        side_border.set(qn('w:sz'), str(size))
        side_border.set(qn('w:space'), '0')
        side_border.set(qn('w:color'), color)

    @staticmethod
    def set_table_row_border(row: _Row, direction: str, size: int, color=BLACK):
        for cell in row.cells:
            DocGenerator.set_table_cell_border(cell, direction, size, color)

    @staticmethod
    def set_table_column_border(column: _Column, direction: str, size: int, color=BLACK):
        for cell in column.cells:
            DocGenerator.set_table_cell_border(cell, direction, size, color)

    @staticmethod
    def set_table_cell_margins(table: Table, top=100, start=100, bottom=100, end=100):
        """
        Set cell padding for all cells in a table (in twips: 1/20 of a point).
        Defaults to 100 twips = 5 points = ~1.76 mm
        """
        table_properties = table._tbl.tblPr

        table_cell_margin = table_properties.find(qn('w:tblCellMar'))
        if table_cell_margin is None:
            table_cell_margin = OxmlElement('w:tblCellMar')
            table_properties.append(table_cell_margin)

        for margin_type, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
            node = table_cell_margin.find(qn(f'w:{margin_type}'))
            if node is None:
                node = OxmlElement(f'w:{margin_type}')
                table_cell_margin.append(node)
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')  # dxa = twentieths of a point

    def generate_doc(self, rechnung_nr: int, rechnung_date: str, price_table_data: List[Tuple[str, str, str]],
                     issue_date: datetime):
        year = issue_date.year

        doc = Document()

        # region global style (text)
        doc_style = doc.styles['Normal']
        doc_style_font = doc_style.font
        doc_style_font.name = 'Times New Roman'
        doc_style_font.size = Pt(12)

        doc_style.paragraph_format.space_before = Pt(0)
        doc_style.paragraph_format.space_after = Pt(0)

        doc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # endregion

        # region section style
        section = doc.sections[0]

        # A4 paper size https://en.wikipedia.org/wiki/ISO_216#A_series
        section.page_height = Mm(297)
        section.page_width = Mm(210)

        margin = Inches(1)
        section.left_margin = margin
        section.right_margin = margin
        section.top_margin = margin
        section.bottom_margin = margin
        # endregion

        # region header table
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row = header_table.rows[0]
        left_cell = row.cells[0]
        right_cell = row.cells[1]

        table_width = section.page_width - section.left_margin - section.right_margin
        left_column_width_ratio = 3 / 5
        left_column_width = int(left_column_width_ratio * table_width)
        header_table.columns[0].width = left_column_width
        header_table.columns[1].width = table_width - left_column_width

        left_cell.text = self.__pr_data[DocGenerator.PrDataKeys.HEADER_TEXT_LEFT]
        right_cell.text = self.__pr_data[DocGenerator.PrDataKeys.HEADER_TEXT_RIGHT]
        # endregion

        # region title (rechnung nr)
        doc.add_paragraph("\n")

        title = doc.add_paragraph()
        title_runner = title.add_run(f"RECHNUNG {rechnung_nr}/{year}")
        title_runner.bold = True

        doc.add_paragraph("\n")
        # endregion

        # region price table
        period_and_billing_text = f"Für den leistungszeitraum von {rechnung_date} erlaube Ich mir Ihnen folgende Leistungen zu berechnen:"
        doc.add_paragraph(period_and_billing_text)

        doc.add_paragraph("")

        price_table = doc.add_table(rows=len(price_table_data), cols=4)
        price_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        price_table.columns[1].width = Inches(1.80)
        price_table.columns[2].width = Inches(0.71)
        price_table.columns[3].width = Inches(0.96)
        price_table.columns[0].width = table_width - sum(map(lambda x: x.width, list(price_table.columns)[1:4]))

        vertical_margin = 80
        horizontal_margin = 120
        table_border_size = 8
        table_group_separator_border_size = int(table_border_size * 2.25)
        for i, row in enumerate(price_table.rows):
            for j, cell in enumerate(row.cells):
                for direction in ('top', 'left'):
                    is_left_of_result_details_column = j == len(row.cells) - 1 and direction == 'left'
                    should_set_border = not is_left_of_result_details_column

                    left_margin = 60 if is_left_of_result_details_column else horizontal_margin
                    right_margin = 0 if is_left_of_result_details_column else horizontal_margin
                    DocGenerator.set_table_cell_margins(price_table, vertical_margin, left_margin, vertical_margin,
                                                        right_margin)

                    if should_set_border:
                        DocGenerator.set_table_cell_border(cell, direction, table_border_size, BLACK)
        DocGenerator.set_table_column_border(price_table.columns[-1], 'right', table_border_size, BLACK)
        DocGenerator.set_table_row_border(price_table.rows[-1], 'bottom', table_border_size, BLACK)

        DocGenerator.set_table_row_border(price_table.rows[-1], 'top', table_group_separator_border_size, BLACK)
        DocGenerator.set_table_row_border(price_table.rows[-3], 'top', table_group_separator_border_size, BLACK)

        # TODO: PRICE_TABLE_COMPUTATION_COLUMN justification
        def split_result_text(the_result_text: str):
            separation_index = the_result_text.index(DECIMAL_SEPARATOR) + 3
            value = the_result_text[:separation_index]
            details = the_result_text[separation_index:]
            return value, details

        for i, (name, computation, result) in enumerate(price_table_data):
            result_value, result_details = split_result_text(result)

            column_index_and_text_pairs = [
                (DocGenerator.PRICE_TABLE_INFO_COLUMN, name),
                (DocGenerator.PRICE_TABLE_COMPUTATION_COLUMN, computation),
                (DocGenerator.PRICE_TABLE_RESULT_VALUE_COLUMN, result_value),
                (DocGenerator.PRICE_TABLE_RESULT_DETAILS_COLUMN, result_details),
            ]

            for column_index, text in column_index_and_text_pairs:
                cell = price_table.cell(i, column_index)
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                run = paragraph.add_run(text.strip())
                if column_index == DocGenerator.PRICE_TABLE_RESULT_VALUE_COLUMN:
                    paragraph.alignment = WD_TABLE_ALIGNMENT.RIGHT
                    run.bold = True
        # endregion

        doc.add_paragraph("")

        # region bank account
        transfer_text = "Bitte überweisen Sie den Betrag innerhalb von 14 Tagen auf dem untenstehenden Konto:"
        doc.add_paragraph(transfer_text)

        doc.add_paragraph("\n")

        bank_account_text_lines = list(
            map(lambda line: f"\t{line}", self.__pr_data[DocGenerator.PrDataKeys.BANK_ACCOUNT].split(os.linesep)))
        bank_account_paragraph = doc.add_paragraph()
        bank_account_paragraph.add_run(bank_account_text_lines[0] + "\n")
        bank_account_paragraph.add_run(bank_account_text_lines[1] + "\n").font.size = doc_style_font.size + Pt(1)
        bank_account_paragraph.add_run(bank_account_text_lines[2] + "\n").bold = True
        bank_account_paragraph.add_run("\n".join(bank_account_text_lines[3:]))

        # endregion

        doc.add_paragraph("\n")

        doc.add_paragraph(get_issue_date_string(issue_date))

        file_name_stem = f"rechnung_{rechnung_nr:03}_{year}"

        DocGenerator.save_doc_and_pdf(doc, file_name_stem)

    @staticmethod
    def save_doc_and_pdf(doc: Document, file_name_stem: str):
        doc_filename = f"{file_name_stem}.docx"
        doc.save(doc_filename)
        try:
            subprocess.run([
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", os.getcwd(),
                doc_filename
            ], check=True)
        except subprocess.CalledProcessError as e:
            print("LibreOffice PDF conversion failed:", e)


def round_half_down(value: Decimal, digits=2):
    quantize_str = '1.' + '0' * digits  # e.g., '1.00' for 2 digits
    return value.quantize(Decimal(quantize_str), rounding=ROUND_HALF_DOWN)


def compute_values(date_list: List[Tuple[str, RechnungInfo]], first_rechnung_nr: int, price_per_stunden: float):
    doc_generator = DocGenerator()
    format_row_string = lambda x, y, z: f"{x:<32}\t{y:>25}\t{z:<25}\n"
    format_computation_column = lambda hours, price: f"{hours:04.2f} St x {price:5.2f} Euro".replace(".",
                                                                                                     DECIMAL_SEPARATOR)
    format_final_price = lambda price: f"{price:7.2f}".replace(".", DECIMAL_SEPARATOR)
    format_price_no_justify = lambda price: f"{price:.2f}".replace(".", DECIMAL_SEPARATOR)
    row_width = 83
    row_group_separator = "-" * row_width + "\n"

    rechnung_prices = []
    issue_date = datetime.now()
    print(get_issue_date_string(issue_date) + "\n")

    for i, (rechnung_date, tuple_list) in enumerate(date_list):
        result = ""

        total_stunden_price = 0
        total_stunden_pl = 0
        total_stunden_pl_price = 0

        price_table_data = []

        for j, (kn_nr, bau, stunden, stunden_pl) in enumerate(tuple_list):
            kn_price = stunden * price_per_stunden
            total_stunden_price += kn_price
            total_stunden_pl += stunden_pl

            bau_string = ' / '.join(map(lambda x: x.strip().capitalize() + 'bau', bau.split('+')))
            info_column = f"{j + 1}. KN NR: {kn_nr} {bau_string}"
            computation_column = format_computation_column(stunden, price_per_stunden)
            result_column = f"{format_final_price(kn_price)} Euro netto"

            result += format_row_string(info_column, computation_column, result_column)
            price_table_data.append((info_column, computation_column, result_column))

        if total_stunden_pl > 0:
            info_column = f'{len(tuple_list) + 1}. Projekt Leiter Zuschlag'
            computation_column = format_computation_column(total_stunden_pl, PRICE_PER_STUNDEN_PL)
            total_stunden_pl_price = total_stunden_pl * PRICE_PER_STUNDEN_PL
            result_column = f"{format_final_price(total_stunden_pl_price)} Euro netto"
            result += format_row_string(info_column, computation_column, result_column)
            price_table_data.append((info_column, computation_column, result_column))

        result += row_group_separator

        arbeitsleistung_netto = total_stunden_price + total_stunden_pl_price
        info_column = "Arbeitsleistung Netto:"
        result_column = f"{format_final_price(arbeitsleistung_netto)} Euro netto"
        result += format_row_string(info_column, "", result_column)
        price_table_data.append((info_column, "", result_column))

        def as_decimal(value: float | int):
            return Decimal(str(value))

        mehrwertsteuer_not_rounded = as_decimal(VAT_RATE) / as_decimal(100) * as_decimal(arbeitsleistung_netto)
        mehrwertsteuer = round_half_down(mehrwertsteuer_not_rounded)
        row_end = f" ({mehrwertsteuer_not_rounded})" if str(mehrwertsteuer_not_rounded).split('.')[1][2] != '0' else ""
        info_column = f"MwSt. {VAT_RATE}%:"
        result_column = f"{format_final_price(mehrwertsteuer)} Euro"
        result += format_row_string(f"MwSt. {VAT_RATE}%:", "", result_column + row_end)
        price_table_data.append((info_column, "", result_column))

        result += row_group_separator

        gesamtbetrag = as_decimal(arbeitsleistung_netto) + mehrwertsteuer
        info_column = "Gesamtbetrag:"
        result_column = f"{format_final_price(gesamtbetrag)} Euro Brutto"
        result += format_row_string(info_column, "", result_column)
        price_table_data.append((info_column, "", result_column))

        rechnung_prices.append(gesamtbetrag)

        rechnung_nr = first_rechnung_nr + i
        print(f"RECHNUNG {rechnung_nr}\t{rechnung_date}")
        print(result)

        doc_generator.generate_doc(rechnung_nr, rechnung_date, price_table_data, issue_date)

        print("\n" * 2)

    total = sum(rechnung_prices)
    total_formatted = format_price_no_justify(total)
    summation = ' + '.join(map(lambda price: f'{format_price_no_justify(price)}', rechnung_prices))
    print(f"total: {summation} = {total_formatted}")

    print()

    last_rechnung_nr = first_rechnung_nr + len(date_list) - 1
    date_range = f"{date_list[0][0]} - {date_list[-1][0]}"
    print(f"email subject: rechnungen {first_rechnung_nr} bis {last_rechnung_nr} | {date_range}")
    print(f"email body: Ins gesamt {total_formatted} Euro.")


def test_round_half_down_2_digit():
    for tenth in range(10):
        for hundredth in range(10):
            for thousandth in range(10):
                number_to_round = Decimal(f"0.{tenth}{hundredth}{thousandth}")
                expected = Decimal(f"0.{tenth}{hundredth}") + Decimal(f"0.0{int(thousandth > 5)}")
                actual = round_half_down(number_to_round)
                assert actual == expected


def main():
    test_round_half_down_2_digit()

    csv_file_path, first_rechnung_nr, price_per_stunden = get_input_args()
    date_list = parse_csv(csv_file_path)
    compute_values(date_list, first_rechnung_nr, price_per_stunden)


if __name__ == "__main__":
    main()
